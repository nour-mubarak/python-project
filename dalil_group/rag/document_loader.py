#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
RAG (Retrieval-Augmented Generation) System
=============================================

Complete production RAG system with document ingestion, retrieval, and generation.
"""

import os
import json
import hashlib
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass, asdict
from datetime import datetime
from abc import ABC, abstractmethod


@dataclass
class Document:
    """Represents a document in the knowledge base."""

    id: str
    title: str
    content: str
    source: str  # 'file', 'url', 'api', 'database'
    chunk_id: Optional[int] = None
    metadata: Dict[str, Any] = None
    embedding: Optional[List[float]] = None
    created_at: str = None

    def __post_init__(self):
        if self.metadata is None:
            self.metadata = {}
        if self.created_at is None:
            self.created_at = datetime.now().isoformat()


@dataclass
class RetrievalResult:
    """Result from document retrieval."""

    document: Document
    score: float  # Similarity score 0-1
    rank: int


class DocumentLoader(ABC):
    """Abstract base for document loaders."""

    @abstractmethod
    def load(self, source: str) -> List[Document]:
        pass


class FileDocumentLoader(DocumentLoader):
    """Load documents from local files."""

    SUPPORTED_EXTENSIONS = {
        ".txt": "text",
        ".md": "markdown",
        ".pdf": "pdf",
        ".docx": "docx",
        ".json": "json",
        ".csv": "csv",
    }

    def load(self, source: str) -> List[Document]:
        """Load documents from file or directory."""
        path = Path(source)

        if path.is_file():
            return [self._load_file(path)]
        elif path.is_dir():
            documents = []
            for file_path in path.rglob("*"):
                if file_path.is_file():
                    doc = self._load_file(file_path)
                    if doc:
                        documents.append(doc)
            return documents
        else:
            raise FileNotFoundError(f"Path not found: {source}")

    def _load_file(self, file_path: Path) -> Optional[Document]:
        """Load a single file."""
        ext = file_path.suffix.lower()

        if ext not in self.SUPPORTED_EXTENSIONS:
            return None

        file_type = self.SUPPORTED_EXTENSIONS[ext]

        try:
            if file_type == "text":
                return self._load_text(file_path)
            elif file_type == "markdown":
                return self._load_markdown(file_path)
            elif file_type == "pdf":
                return self._load_pdf(file_path)
            elif file_type == "docx":
                return self._load_docx(file_path)
            elif file_type == "json":
                return self._load_json(file_path)
            elif file_type == "csv":
                return self._load_csv(file_path)
        except Exception as e:
            print(f"Error loading {file_path}: {e}")
            return None

    def _load_text(self, file_path: Path) -> Document:
        """Load plain text file."""
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()

        return Document(
            id=self._generate_id(file_path),
            title=file_path.stem,
            content=content,
            source="file",
            metadata={"path": str(file_path), "format": "text"},
        )

    def _load_markdown(self, file_path: Path) -> Document:
        """Load markdown file."""
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()

        return Document(
            id=self._generate_id(file_path),
            title=file_path.stem,
            content=content,
            source="file",
            metadata={"path": str(file_path), "format": "markdown"},
        )

    def _load_pdf(self, file_path: Path) -> Optional[Document]:
        """Load PDF file."""
        try:
            import PyPDF2

            text = ""
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text += page.extract_text()

            return Document(
                id=self._generate_id(file_path),
                title=file_path.stem,
                content=text,
                source="file",
                metadata={"path": str(file_path), "format": "pdf"},
            )
        except ImportError:
            print("PyPDF2 not installed: pip install PyPDF2")
            return None

    def _load_docx(self, file_path: Path) -> Optional[Document]:
        """Load DOCX file."""
        try:
            from docx import Document as DocxDocument

            doc = DocxDocument(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])

            return Document(
                id=self._generate_id(file_path),
                title=file_path.stem,
                content=text,
                source="file",
                metadata={"path": str(file_path), "format": "docx"},
            )
        except ImportError:
            print("python-docx not installed: pip install python-docx")
            return None

    def _load_json(self, file_path: Path) -> Document:
        """Load JSON file."""
        with open(file_path, "r", encoding="utf-8") as f:
            data = json.load(f)

        # Convert JSON to readable text
        content = json.dumps(data, indent=2, ensure_ascii=False)

        return Document(
            id=self._generate_id(file_path),
            title=file_path.stem,
            content=content,
            source="file",
            metadata={"path": str(file_path), "format": "json"},
        )

    def _load_csv(self, file_path: Path) -> Optional[Document]:
        """Load CSV file."""
        try:
            import csv

            rows = []
            with open(file_path, "r", encoding="utf-8") as f:
                reader = csv.DictReader(f)
                for row in reader:
                    rows.append(row)

            content = json.dumps(rows, indent=2, ensure_ascii=False)

            return Document(
                id=self._generate_id(file_path),
                title=file_path.stem,
                content=content,
                source="file",
                metadata={"path": str(file_path), "format": "csv"},
            )
        except Exception as e:
            print(f"Error loading CSV: {e}")
            return None

    @staticmethod
    def _generate_id(file_path: Path) -> str:
        """Generate unique ID from file path."""
        path_str = str(file_path.absolute())
        return hashlib.md5(path_str.encode()).hexdigest()


class URLDocumentLoader(DocumentLoader):
    """Load documents from web URLs."""

    def load(self, source: str) -> List[Document]:
        """Load document from URL."""
        try:
            import requests
            from bs4 import BeautifulSoup

            print(f"Fetching: {source}")
            response = requests.get(source, timeout=10)
            response.raise_for_status()

            # Parse HTML
            soup = BeautifulSoup(response.text, "html.parser")

            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()

            # Get text
            text = soup.get_text(separator="\n")

            # Clean up text
            lines = [line.strip() for line in text.split("\n") if line.strip()]
            content = "\n".join(lines)

            doc = Document(
                id=hashlib.md5(source.encode()).hexdigest(),
                title=soup.title.string or source,
                content=content,
                source="url",
                metadata={"url": source},
            )

            return [doc]

        except ImportError:
            print(
                "requests/beautifulsoup4 not installed: pip install requests beautifulsoup4"
            )
            return []
        except Exception as e:
            print(f"Error loading URL {source}: {e}")
            return []


class DocumentChunker:
    """Split documents into chunks for embedding."""

    @staticmethod
    def chunk_by_tokens(
        document: Document,
        chunk_size: int = 512,
        overlap: int = 50,
    ) -> List[Document]:
        """
        Split document into chunks by approximate token count.

        Assumes average 4 characters per token.
        """
        chunks = []
        content = document.content

        # Rough token estimation: 1 token ≈ 4 characters
        char_per_token = 4
        char_chunk_size = chunk_size * char_per_token
        char_overlap = overlap * char_per_token

        start = 0
        chunk_id = 0

        while start < len(content):
            end = min(start + char_chunk_size, len(content))

            # Try to break at sentence boundary
            if end < len(content):
                # Find last sentence break
                last_period = content.rfind(".", start, end)
                if last_period > start + char_chunk_size // 2:
                    end = last_period + 1

            chunk_text = content[start:end].strip()

            if chunk_text:
                chunk_doc = Document(
                    id=f"{document.id}_chunk_{chunk_id}",
                    title=document.title,
                    content=chunk_text,
                    source=document.source,
                    chunk_id=chunk_id,
                    metadata={
                        **document.metadata,
                        "parent_id": document.id,
                        "chunk_position": chunk_id,
                    },
                )
                chunks.append(chunk_doc)
                chunk_id += 1

            start = end - char_overlap

        return chunks if chunks else [document]

    @staticmethod
    def chunk_by_paragraphs(
        document: Document,
        max_chunk_size: int = 512,
    ) -> List[Document]:
        """Split document by paragraphs."""
        chunks = []
        paragraphs = document.content.split("\n\n")

        current_chunk = ""
        chunk_id = 0

        for para in paragraphs:
            if len(current_chunk) + len(para) < max_chunk_size * 4:
                current_chunk += para + "\n\n"
            else:
                if current_chunk.strip():
                    chunk_doc = Document(
                        id=f"{document.id}_chunk_{chunk_id}",
                        title=document.title,
                        content=current_chunk.strip(),
                        source=document.source,
                        chunk_id=chunk_id,
                        metadata={
                            **document.metadata,
                            "parent_id": document.id,
                        },
                    )
                    chunks.append(chunk_doc)
                    chunk_id += 1

                current_chunk = para + "\n\n"

        if current_chunk.strip():
            chunk_doc = Document(
                id=f"{document.id}_chunk_{chunk_id}",
                title=document.title,
                content=current_chunk.strip(),
                source=document.source,
                chunk_id=chunk_id,
                metadata={
                    **document.metadata,
                    "parent_id": document.id,
                },
            )
            chunks.append(chunk_doc)

        return chunks if chunks else [document]
