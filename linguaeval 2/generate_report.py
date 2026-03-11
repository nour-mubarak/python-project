#!/usr/bin/env python3
"""
LinguaEval — Report Generator
================================

Generates a professional Multilingual AI Readiness Report (DOCX)
from evaluation results JSON.

Usage:
    python generate_report.py --results results/client_results.json
    python generate_report.py --results results/client_results.json --output reports/client_report.docx
"""

import argparse
import json
import os
import sys
from datetime import datetime
from typing import Dict, List, Optional

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ═══════════════════════════════════════════════════════════════
# DESIGN CONSTANTS
# ═══════════════════════════════════════════════════════════════

NAVY = RGBColor(0x0C, 0x2D, 0x48)
ACCENT = RGBColor(0x02, 0x84, 0xC7)
DARK = RGBColor(0x1E, 0x29, 0x3B)
GRAY = RGBColor(0x64, 0x74, 0x8B)
GREEN = RGBColor(0x16, 0x65, 0x34)
RED = RGBColor(0xB9, 0x1C, 0x1C)
ORANGE = RGBColor(0xC2, 0x41, 0x0C)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BG = "EBF5FB"

FONT_BODY = "Calibri"
FONT_SIZE = Pt(10.5)


def severity_color(severity: str) -> RGBColor:
    """Get colour for severity level."""
    return {
        "low": GREEN,
        "medium": ORANGE,
        "high": RED,
        "critical": RED,
    }.get(severity, GRAY)


def score_to_status(score: float) -> str:
    """Convert average score to deployment status."""
    if score >= 80:
        return "Ready for Pilot"
    elif score >= 65:
        return "Restricted Pilot Only"
    return "Not Ready"


def score_to_color(score: float) -> RGBColor:
    if score >= 80:
        return GREEN
    elif score >= 65:
        return ORANGE
    return RED


# ═══════════════════════════════════════════════════════════════
# DOCUMENT HELPERS
# ═══════════════════════════════════════════════════════════════

def set_cell_shading(cell, color_hex: str):
    """Set background colour on a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_paragraph(doc, text, style="Normal", bold=False, color=None,
                         size=None, alignment=None, space_after=None, space_before=None):
    """Add a paragraph with custom styling."""
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    
    run = p.add_run(text)
    run.font.name = FONT_BODY
    run.font.size = size or FONT_SIZE
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    return p


def add_rich_paragraph(doc, runs_data, alignment=None, space_after=None):
    """Add a paragraph with mixed formatting."""
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    
    for rd in runs_data:
        run = p.add_run(rd.get("text", ""))
        run.font.name = FONT_BODY
        run.font.size = rd.get("size", FONT_SIZE)
        run.font.bold = rd.get("bold", False)
        run.font.italic = rd.get("italic", False)
        if rd.get("color"):
            run.font.color.rgb = rd["color"]
    return p


def add_key_value_table(doc, data: List[tuple], col_widths=(Inches(2), Inches(4.5))):
    """Add a simple two-column key-value table."""
    table = doc.add_table(rows=len(data), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for i, (key, value) in enumerate(data):
        # Key cell
        cell_k = table.cell(i, 0)
        cell_k.width = col_widths[0]
        p = cell_k.paragraphs[0]
        run = p.add_run(key)
        run.font.name = FONT_BODY
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = NAVY
        set_cell_shading(cell_k, "F8FAFC")
        
        # Value cell
        cell_v = table.cell(i, 1)
        cell_v.width = col_widths[1]
        p = cell_v.paragraphs[0]
        run = p.add_run(str(value))
        run.font.name = FONT_BODY
        run.font.size = Pt(10)
        run.font.color.rgb = DARK
    
    return table


def add_scores_table(doc, models_data: dict, dimension_names: list):
    """Add a model comparison scores table."""
    models = list(models_data.keys())
    n_cols = 1 + len(models) * 2 + len(models)  # dim + (en + ar) per model + gap per model
    
    # Simplified: dim | model1_en | model1_ar | model1_gap | model2_en | ...
    header_cols = ["Dimension"]
    for m in models:
        short_name = m.split("-")[0].upper() if "-" in m else m.upper()
        header_cols.extend([f"{short_name} EN", f"{short_name} AR", f"{short_name} Gap"])
    
    table = doc.add_table(rows=1 + len(dimension_names), cols=len(header_cols))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header row
    for j, header_text in enumerate(header_cols):
        cell = table.cell(0, j)
        p = cell.paragraphs[0]
        run = p.add_run(header_text)
        run.font.name = FONT_BODY
        run.font.size = Pt(8.5)
        run.font.bold = True
        run.font.color.rgb = WHITE
        set_cell_shading(cell, "0C2D48")
    
    # Data rows
    for i, dim in enumerate(dimension_names):
        row_idx = i + 1
        # Dimension name
        cell = table.cell(row_idx, 0)
        p = cell.paragraphs[0]
        run = p.add_run(dim.replace("_", " ").title())
        run.font.name = FONT_BODY
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = DARK
        
        if i % 2 == 0:
            set_cell_shading(cell, "F8FAFC")
        
        col_offset = 1
        for m in models:
            m_data = models_data.get(m, {})
            en_score = m_data.get("en", {}).get(dim, {}).get("average", "N/A")
            ar_score = m_data.get("ar", {}).get(dim, {}).get("average", "N/A")
            gap_data = m_data.get("cross_lingual_gap", {}).get(dim, {})
            gap = gap_data.get("gap", "N/A")
            
            for k, (val, is_gap) in enumerate([(en_score, False), (ar_score, False), (gap, True)]):
                cell = table.cell(row_idx, col_offset + k)
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                if isinstance(val, (int, float)):
                    text = f"{val:.1f}%"
                    run = p.add_run(text)
                    run.font.name = FONT_BODY
                    run.font.size = Pt(9)
                    run.font.bold = True
                    if is_gap:
                        sev = gap_data.get("severity", "low")
                        run.font.color.rgb = severity_color(sev)
                    else:
                        run.font.color.rgb = score_to_color(val)
                else:
                    run = p.add_run("N/A")
                    run.font.name = FONT_BODY
                    run.font.size = Pt(9)
                    run.font.color.rgb = GRAY
                
                if i % 2 == 0:
                    set_cell_shading(cell, "F8FAFC")
            
            col_offset += 3
    
    return table


# ═══════════════════════════════════════════════════════════════
# REPORT GENERATOR
# ═══════════════════════════════════════════════════════════════

class ReportGenerator:
    """Generates the Multilingual AI Readiness Report from evaluation results."""

    def __init__(self, results_path: str):
        """Load results from JSON file."""
        with open(results_path, encoding="utf-8") as f:
            self.data = json.load(f)
        
        self.metadata = self.data.get("metadata", {})
        self.aggregates = self.data.get("aggregates", {})
        self.detailed = self.data.get("detailed_results", [])
        
        self.client_name = self.metadata.get("client_name", "Client")
        self.sector = self.metadata.get("sector", "general")
        self.models = self.metadata.get("models", [])
        self.timestamp = self.metadata.get("timestamp", datetime.now().isoformat())

    def _calculate_overall_scores(self) -> dict:
        """Calculate overall scores per model."""
        overall = {}
        for model_id, data in self.aggregates.items():
            en_scores = [v["average"] for v in data.get("en", {}).values()]
            ar_scores = [v["average"] for v in data.get("ar", {}).values()]
            gaps = [v["gap"] for v in data.get("cross_lingual_gap", {}).values()]
            
            en_avg = sum(en_scores) / len(en_scores) if en_scores else 0
            ar_avg = sum(ar_scores) / len(ar_scores) if ar_scores else 0
            avg_gap = sum(gaps) / len(gaps) if gaps else 0
            combined = (en_avg + ar_avg) / 2
            
            overall[model_id] = {
                "en_average": round(en_avg, 1),
                "ar_average": round(ar_avg, 1),
                "combined_average": round(combined, 1),
                "avg_gap": round(avg_gap, 1),
                "status": score_to_status(combined),
            }
        return overall

    def _collect_flags(self, model_id: str = None, dimension: str = None, language: str = None) -> List[dict]:
        """Collect all flags from detailed results, optionally filtered."""
        flags = []
        for result in self.detailed:
            for mid, model_data in result.get("model_scores", {}).items():
                if model_id and mid != model_id:
                    continue
                for lang, lang_data in model_data.items():
                    if language and lang != language:
                        continue
                    for score_entry in lang_data.get("scores", []):
                        if dimension and score_entry["dimension"] != dimension:
                            continue
                        for flag in score_entry.get("flags", []):
                            flags.append({
                                "prompt_id": result["prompt_id"],
                                "model": mid,
                                "language": lang,
                                "dimension": score_entry["dimension"],
                                "severity": score_entry["severity"],
                                "flag": flag,
                                "prompt_en": result.get("prompt_en", ""),
                                "prompt_ar": result.get("prompt_ar", ""),
                                "response": lang_data.get("response", "")[:200],
                            })
        return flags

    def generate(self, output_path: str):
        """Generate the complete Readiness Report."""
        doc = Document()
        
        # ── Page setup ──
        section = doc.sections[0]
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        
        # ── Default styles ──
        style = doc.styles["Normal"]
        style.font.name = FONT_BODY
        style.font.size = FONT_SIZE
        style.font.color.rgb = DARK
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.15
        
        # Configure heading styles
        for level, (size, color) in enumerate([(Pt(22), NAVY), (Pt(16), ACCENT), (Pt(13), DARK)], 1):
            heading_style = doc.styles[f"Heading {level}"]
            heading_style.font.name = FONT_BODY
            heading_style.font.size = size
            heading_style.font.color.rgb = color
            heading_style.font.bold = True
            heading_style.paragraph_format.space_before = Pt(18 if level == 1 else 14)
            heading_style.paragraph_format.space_after = Pt(8)

        overall = self._calculate_overall_scores()

        # ═══════════════════════════════════════
        # COVER PAGE
        # ═══════════════════════════════════════
        for _ in range(6):
            doc.add_paragraph()
        
        add_styled_paragraph(doc, "MULTILINGUAL AI", bold=True, color=NAVY,
                           size=Pt(28), alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
        add_styled_paragraph(doc, "READINESS REPORT", bold=True, color=NAVY,
                           size=Pt(28), alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=16)
        
        add_styled_paragraph(doc, f"Prepared for {self.client_name}", color=ACCENT,
                           size=Pt(14), alignment=WD_ALIGN_PARAGRAPH.CENTER,
                           bold=True, space_after=24)
        
        # Divider
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("─" * 40)
        run.font.color.rgb = ACCENT
        run.font.size = Pt(10)
        
        date_str = datetime.fromisoformat(self.timestamp).strftime("%B %Y")
        add_styled_paragraph(doc, date_str, color=GRAY, size=Pt(12),
                           alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
        add_styled_paragraph(doc, "LinguaEval", color=GRAY, size=Pt(12),
                           alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Confidential")
        run.font.name = FONT_BODY
        run.font.size = Pt(10)
        run.font.color.rgb = GRAY
        run.font.italic = True
        
        doc.add_page_break()

        # ═══════════════════════════════════════
        # EXECUTIVE SUMMARY
        # ═══════════════════════════════════════
        doc.add_heading("1. Executive Summary", level=1)
        
        doc.add_paragraph(
            f"This report presents the results of a Multilingual AI Readiness Assessment "
            f"conducted for {self.client_name}. The evaluation tested {len(self.models)} AI model(s) "
            f"across Arabic and English on {self.metadata.get('total_prompts', 'N/A')} prompts "
            f"covering {self.sector} use cases."
        )
        
        # Overall status per model
        doc.add_heading("Deployment Readiness Summary", level=2)
        
        status_data = []
        for model_id, scores in overall.items():
            short = model_id.split("-")[0].upper() if "-" in model_id else model_id.upper()
            status_data.append((
                short,
                f"{scores['en_average']}%",
                f"{scores['ar_average']}%",
                f"{scores['avg_gap']}%",
                scores['status'],
            ))
        
        table = doc.add_table(rows=1 + len(status_data), cols=5)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        headers = ["Model", "English Score", "Arabic Score", "Avg Gap", "Readiness Status"]
        for j, h in enumerate(headers):
            cell = table.cell(0, j)
            p = cell.paragraphs[0]
            run = p.add_run(h)
            run.font.name = FONT_BODY
            run.font.size = Pt(9)
            run.font.bold = True
            run.font.color.rgb = WHITE
            set_cell_shading(cell, "0C2D48")
        
        for i, row_data in enumerate(status_data):
            for j, val in enumerate(row_data):
                cell = table.cell(i + 1, j)
                p = cell.paragraphs[0]
                run = p.add_run(str(val))
                run.font.name = FONT_BODY
                run.font.size = Pt(9.5)
                if j == 4:  # Status column
                    run.font.bold = True
                    if "Ready for Pilot" == val:
                        run.font.color.rgb = GREEN
                    elif "Restricted" in val:
                        run.font.color.rgb = ORANGE
                    else:
                        run.font.color.rgb = RED
                elif j in (1, 2):
                    score_val = float(val.replace("%", ""))
                    run.font.color.rgb = score_to_color(score_val)
                    run.font.bold = True
        
        doc.add_paragraph()
        
        # Key finding
        best_model = max(overall.items(), key=lambda x: x[1]["combined_average"])
        doc.add_paragraph(
            f"Based on this evaluation, {best_model[0]} achieved the highest combined score "
            f"of {best_model[1]['combined_average']}% across both languages. "
            f"Its deployment readiness status is: {best_model[1]['status']}."
        )
        
        doc.add_page_break()

        # ═══════════════════════════════════════
        # METHODOLOGY
        # ═══════════════════════════════════════
        doc.add_heading("2. Methodology", level=1)
        
        doc.add_paragraph(
            "This assessment used the LinguaEval evaluation framework, which tests AI models "
            "across six dimensions in both Arabic and English using sector-specific prompt packs."
        )
        
        doc.add_heading("Evaluation Parameters", level=2)
        add_key_value_table(doc, [
            ("Client", self.client_name),
            ("Sector", self.sector.title()),
            ("Use Case", self.metadata.get("use_case", "General")),
            ("Models Evaluated", ", ".join(self.models)),
            ("Languages", "English, Arabic"),
            ("Total Prompts", str(self.metadata.get("total_prompts", "N/A"))),
            ("Evaluation Date", datetime.fromisoformat(self.timestamp).strftime("%d %B %Y")),
            ("Privacy Mode", self.metadata.get("privacy_mode", "mode_a").replace("_", " ").title()),
        ])
        
        doc.add_paragraph()
        
        doc.add_heading("Evaluation Dimensions", level=2)
        dimensions_desc = [
            ("Factual Accuracy", "Verifies responses against known ground truth and key facts."),
            ("Gender Bias", "Detects gendered defaults, stereotypical language, and bias asymmetries across languages."),
            ("Hallucination", "Identifies fabricated facts, unverified claims, and contradictions."),
            ("Cross-Lingual Consistency", "Measures whether Arabic and English responses provide equivalent information."),
            ("Cultural Sensitivity", "Evaluates cultural appropriateness and avoidance of stereotypes."),
            ("Fluency & Coherence", "Assesses language quality, completeness, and code-switching."),
        ]
        
        for dim_name, dim_desc in dimensions_desc:
            p = doc.add_paragraph()
            run = p.add_run(f"{dim_name}: ")
            run.font.bold = True
            run.font.color.rgb = NAVY
            run.font.size = Pt(10)
            run = p.add_run(dim_desc)
            run.font.size = Pt(10)

        doc.add_page_break()

        # ═══════════════════════════════════════
        # DETAILED RESULTS
        # ═══════════════════════════════════════
        doc.add_heading("3. Detailed Results", level=1)
        
        # Get all unique dimensions from the data
        all_dims = set()
        for model_data in self.aggregates.values():
            all_dims.update(model_data.get("en", {}).keys())
            all_dims.update(model_data.get("ar", {}).keys())
        all_dims = sorted(all_dims)
        
        if all_dims and self.aggregates:
            doc.add_heading("3.1 Model Comparison", level=2)
            doc.add_paragraph(
                "The following table shows average scores per dimension for each model, "
                "comparing English and Arabic performance with the cross-lingual gap."
            )
            add_scores_table(doc, self.aggregates, all_dims)
            doc.add_paragraph()

        # ── Per-model analysis ──
        for idx, (model_id, model_data) in enumerate(self.aggregates.items()):
            model_scores = overall.get(model_id, {})
            short_name = model_id.split("-")[0].upper() if "-" in model_id else model_id.upper()
            
            doc.add_heading(f"3.{idx + 2} {short_name} Analysis", level=2)
            
            add_rich_paragraph(doc, [
                {"text": "Overall English Score: ", "bold": True, "color": GRAY, "size": Pt(10)},
                {"text": f"{model_scores.get('en_average', 'N/A')}%", "bold": True,
                 "color": score_to_color(model_scores.get('en_average', 0)), "size": Pt(11)},
                {"text": "    Overall Arabic Score: ", "bold": True, "color": GRAY, "size": Pt(10)},
                {"text": f"{model_scores.get('ar_average', 'N/A')}%", "bold": True,
                 "color": score_to_color(model_scores.get('ar_average', 0)), "size": Pt(11)},
                {"text": "    Average Gap: ", "bold": True, "color": GRAY, "size": Pt(10)},
                {"text": f"{model_scores.get('avg_gap', 'N/A')}%", "bold": True,
                 "color": ORANGE if model_scores.get('avg_gap', 0) > 8 else GREEN, "size": Pt(11)},
            ])
            
            # Collect flags for this model
            model_flags = self._collect_flags(model_id=model_id)
            
            if model_flags:
                doc.add_heading("Flagged Issues", level=3)
                
                # Group flags by dimension
                by_dimension = {}
                for f in model_flags:
                    dim = f["dimension"]
                    if dim not in by_dimension:
                        by_dimension[dim] = []
                    by_dimension[dim].append(f)
                
                for dim, dim_flags in by_dimension.items():
                    doc.add_heading(f"{dim.replace('_', ' ').title()} Flags", level=3)
                    
                    # Show top 5 flags per dimension
                    for flag_data in dim_flags[:5]:
                        p = doc.add_paragraph(style="List Bullet")
                        
                        lang_label = "AR" if flag_data["language"] == "ar" else "EN"
                        run = p.add_run(f"[{lang_label}] [{flag_data['severity'].upper()}] ")
                        run.font.size = Pt(9)
                        run.font.bold = True
                        run.font.color.rgb = severity_color(flag_data["severity"])
                        
                        run = p.add_run(flag_data["flag"])
                        run.font.size = Pt(9)
                        run.font.color.rgb = DARK
                    
                    if len(dim_flags) > 5:
                        add_styled_paragraph(
                            doc,
                            f"... and {len(dim_flags) - 5} additional {dim} flag(s)",
                            color=GRAY, size=Pt(9)
                        )

        doc.add_page_break()

        # ═══════════════════════════════════════
        # CROSS-LINGUAL GAP ANALYSIS
        # ═══════════════════════════════════════
        doc.add_heading("4. Cross-Lingual Gap Analysis", level=1)
        
        doc.add_paragraph(
            "Cross-lingual gaps measure how differently each model performs between English and Arabic. "
            "Larger gaps indicate higher risk for bilingual deployments. Gaps above 12% are flagged as "
            "high-risk and may require language-specific guardrails."
        )
        
        for model_id, model_data in self.aggregates.items():
            short_name = model_id.split("-")[0].upper() if "-" in model_id else model_id.upper()
            gap_data = model_data.get("cross_lingual_gap", {})
            
            if gap_data:
                doc.add_heading(f"{short_name} Cross-Lingual Gaps", level=2)
                
                gap_rows = []
                for dim, gd in gap_data.items():
                    gap_rows.append((
                        dim.replace("_", " ").title(),
                        f"{gd.get('en_avg', 'N/A')}%",
                        f"{gd.get('ar_avg', 'N/A')}%",
                        f"{gd.get('gap', 'N/A')}%",
                        gd.get("severity", "N/A").title(),
                    ))
                
                table = doc.add_table(rows=1 + len(gap_rows), cols=5)
                table.style = "Table Grid"
                
                for j, h in enumerate(["Dimension", "English", "Arabic", "Gap", "Severity"]):
                    cell = table.cell(0, j)
                    run = cell.paragraphs[0].add_run(h)
                    run.font.name = FONT_BODY
                    run.font.size = Pt(9)
                    run.font.bold = True
                    run.font.color.rgb = WHITE
                    set_cell_shading(cell, "0C2D48")
                
                for i, row in enumerate(gap_rows):
                    for j, val in enumerate(row):
                        cell = table.cell(i + 1, j)
                        run = cell.paragraphs[0].add_run(str(val))
                        run.font.name = FONT_BODY
                        run.font.size = Pt(9)
                        if j == 4:  # Severity
                            run.font.bold = True
                            run.font.color.rgb = severity_color(val.lower())
                
                doc.add_paragraph()

        doc.add_page_break()

        # ═══════════════════════════════════════
        # RECOMMENDATIONS
        # ═══════════════════════════════════════
        doc.add_heading("5. Deployment Recommendation", level=1)
        
        for model_id, scores in overall.items():
            short_name = model_id.split("-")[0].upper() if "-" in model_id else model_id.upper()
            status = scores["status"]
            
            doc.add_heading(f"{short_name}: {status}", level=2)
            
            if status == "Ready for Pilot":
                doc.add_paragraph(
                    f"{short_name} demonstrates acceptable performance across both languages for a bounded "
                    f"pilot deployment. Standard monitoring and periodic re-evaluation are recommended."
                )
                doc.add_paragraph("Recommended controls:")
                for ctrl in [
                    "Standard human review sampling (10-15% of outputs)",
                    "Monthly performance monitoring",
                    "Quarterly cross-lingual re-evaluation",
                    "User feedback collection mechanism",
                ]:
                    doc.add_paragraph(ctrl, style="List Bullet")
                    
            elif status == "Restricted Pilot Only":
                doc.add_paragraph(
                    f"{short_name} shows moderate cross-lingual performance gaps that require additional "
                    f"controls before deployment. A restricted pilot with enhanced oversight is recommended."
                )
                doc.add_paragraph("Required controls before deployment:")
                for ctrl in [
                    "Enhanced human review for all Arabic outputs (100% review initially)",
                    "Domain-specific prompt engineering to reduce bias patterns",
                    "Restricted scope: deploy in one workflow only before expanding",
                    "Weekly performance monitoring during pilot phase",
                    "Re-evaluation after 4 weeks with updated scoring",
                ]:
                    doc.add_paragraph(ctrl, style="List Bullet")
                    
            else:
                doc.add_paragraph(
                    f"{short_name} exhibits significant performance issues that make it unsuitable for "
                    f"bilingual deployment without substantial mitigation. Deployment is not recommended "
                    f"at this time."
                )
                doc.add_paragraph("Recommended next steps:")
                for ctrl in [
                    "Evaluate alternative model providers",
                    "Consider Arabic-specialist models (e.g. Jais, Allam) for Arabic-facing applications",
                    "Implement language-specific routing: use different models for Arabic vs English",
                    "Re-evaluate after provider updates or fine-tuning",
                ]:
                    doc.add_paragraph(ctrl, style="List Bullet")

        doc.add_page_break()

        # ═══════════════════════════════════════
        # NEXT STEPS
        # ═══════════════════════════════════════
        doc.add_heading("6. Recommended Next Steps", level=1)
        
        steps = [
            ("Immediate (This Week)", [
                "Review this report with your AI governance or technology leadership team",
                "Identify the highest-risk findings and determine which require immediate action",
                "Decide on deployment status for each evaluated model",
            ]),
            ("Short-Term (Next 30 Days)", [
                "Implement recommended controls for any models approved for pilot",
                "Begin restricted pilot deployment if applicable",
                "Establish monitoring and feedback collection processes",
            ]),
            ("Medium-Term (60-90 Days)", [
                "Conduct follow-up evaluation to measure improvement",
                "Expand pilot scope if initial results are positive",
                "Consider a Cross-Lingual Bias & Reliability Audit for deeper analysis",
                "Evaluate additional models or providers as needed",
            ]),
        ]
        
        for period, actions in steps:
            doc.add_heading(period, level=2)
            for action in actions:
                doc.add_paragraph(action, style="List Bullet")

        # ═══════════════════════════════════════
        # ABOUT / CONTACT
        # ═══════════════════════════════════════
        doc.add_page_break()
        doc.add_heading("About LinguaEval", level=1)
        
        doc.add_paragraph(
            "LinguaEval is a specialist multilingual AI evaluation and deployment studio based in the UK, "
            "serving organisations across the UK and GCC. We help organisations evaluate, de-risk, and deploy "
            "Arabic-English AI systems through research-grade bias auditing, cross-lingual benchmarking, "
            "and high-trust pilot delivery."
        )
        
        doc.add_paragraph()
        
        doc.add_heading("Our Services", level=2)
        services = [
            "Multilingual AI Readiness Assessment — Decision-ready evaluation with deployment recommendation",
            "Cross-Lingual Bias & Reliability Audit — Deep technical analysis with mitigation guidance",
            "High-Trust AI Pilot — Bounded deployment with governance and handover",
            "Assurance & Monitoring Retainer — Ongoing evaluation and advisory",
        ]
        for s in services:
            doc.add_paragraph(s, style="List Bullet")
        
        doc.add_paragraph()
        add_styled_paragraph(doc, "Contact: hello@linguaeval.com", color=ACCENT, bold=True)
        add_styled_paragraph(doc, "Website: www.linguaeval.com", color=ACCENT, bold=True)
        
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run(
            "\u00A9 2026 LinguaEval Ltd. All rights reserved. This report is confidential and prepared "
            "exclusively for the named client. It may not be distributed without written permission."
        )
        run.font.name = FONT_BODY
        run.font.size = Pt(9)
        run.font.color.rgb = GRAY
        run.font.italic = True

        # ── Save ──
        os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else ".", exist_ok=True)
        doc.save(output_path)
        print(f"\n[+] Report generated: {output_path}")
        print(f"    Client: {self.client_name}")
        print(f"    Models: {', '.join(self.models)}")
        print(f"    Prompts: {self.metadata.get('total_prompts', 'N/A')}")


def main():
    parser = argparse.ArgumentParser(description="LinguaEval Report Generator")
    parser.add_argument("--results", required=True, help="Path to evaluation results JSON")
    parser.add_argument("--output", default=None, help="Output path for DOCX report")
    
    args = parser.parse_args()
    
    if not os.path.exists(args.results):
        print(f"[!] Results file not found: {args.results}")
        sys.exit(1)
    
    if args.output is None:
        base = os.path.splitext(os.path.basename(args.results))[0]
        args.output = os.path.join("reports", f"{base}_report.docx")
    
    generator = ReportGenerator(args.results)
    generator.generate(args.output)


if __name__ == "__main__":
    main()
